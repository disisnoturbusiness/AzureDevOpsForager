#!/usr/bin/env python3
"""
build_word_docs.py — Convert the Azure DevOps Forager Markdown guides into
polished Microsoft Word (.docx) documents in a clean, professional corporate
visual style.

Sources -> Outputs (all inside this docs/ folder):
    USER_GUIDE.md       -> AzureDevOpsForager_User_Guide.docx
    DEVELOPER_GUIDE.md  -> AzureDevOpsForager_Developer_Guide.docx

Re-run this script whenever the guides change:
    python build_word_docs.py

Only depends on python-docx (installed).  It is a small, self-contained
Markdown -> docx converter covering the constructs actually used in the two
guides: headings (# .. ####), paragraphs with **bold** / `inline code` /
[links](url), fenced ``` code blocks, pipe tables, bullet + numbered lists
(with nesting), blockquote callouts (>), and horizontal rules (---).
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------- #
# Palette / style constants  (corporate blue palette)
# --------------------------------------------------------------------------- #
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
LAST_UPDATED = "2026"  # fixed placeholder per spec (no datetime.now)

DARK_BLUE = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6  H1/H2/H4, table header fill, cover title
LIGHT_BLUE = RGBColor(0x5B, 0x9B, 0xD5)  # #5B9BD5  H3
GRAY_SUB = RGBColor(0x66, 0x66, 0x66)    # #666666  cover subtitle
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CODE_FILL = "F2F2F2"        # fenced code block background
INLINE_CODE_FILL = "F2F2F2" # inline code very-light shading
CALLOUT_BLUE = "DEEBF7"     # default callout background
CALLOUT_RED = "FBE4E4"      # warning/important callout background
CALLOUT_BORDER_BLUE = "9DC3E6"
CALLOUT_BORDER_RED = "E6A0A0"
CODE_BORDER = "CCCCCC"
TABLE_GRID = "D9D9D9"

BODY_FONT = "Arial"
CODE_FONT = "Consolas"

# Inline-markdown tokenizer: bold (**...**), italic (*...*), inline code (`...`),
# link ([t](u)).  Bold is matched before italic so ** wins over *.
_INLINE_RE = re.compile(
    r"(\*\*.+?\*\*)"              # bold
    r"|(\*(?!\s)[^*]+?(?<!\s)\*)" # italic (single *, no surrounding space)
    r"|(`[^`]+`)"                 # inline code
    r"|(\[[^\]]+\]\([^)]+\))"     # link
)


# --------------------------------------------------------------------------- #
# Low-level OOXML shading / border helpers
# --------------------------------------------------------------------------- #
def _set_shading(element, fill_hex):
    """Apply a solid fill to a paragraph (<w:pPr>) or table cell (<w:tcPr>)."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def _cell_shading(cell, fill_hex):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def _para_shading(paragraph, fill_hex):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    ppr.append(shd)


def _para_borders(paragraph, color_hex, sz=6, space=8, sides=("top", "bottom", "left", "right")):
    """Add a thin box border around a paragraph."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color_hex)
        pbdr.append(el)
    ppr.append(pbdr)


def _cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{name}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tcpr.append(mar)


def _set_table_borders(table, color_hex):
    tbl = table._tbl
    tblpr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tblpr.append(borders)


def _keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    ppr.append(kwn)


# --------------------------------------------------------------------------- #
# Inline run rendering
# --------------------------------------------------------------------------- #
def _add_inline_code_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)
    # very light shading behind inline code
    rpr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), INLINE_CODE_FILL)
    rpr.append(shd)
    return run


def add_inline_markdown(paragraph, text, base_bold=False, base_italic=False, base_color=None):
    """Render text with **bold**, `inline code`, and [link](url) into a paragraph."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _plain_run(paragraph, text[pos:m.start()], base_bold, base_italic, base_color)
        token = m.group(0)
        if token.startswith("**"):
            # bold span may itself contain inline code / a link — recurse so `code` and links
            # inside **...** render correctly instead of leaking their markdown characters.
            add_inline_markdown(paragraph, token[2:-2], base_bold=True,
                                base_italic=base_italic, base_color=base_color)
        elif token.startswith("*"):
            # italic span may itself contain a link / code — recurse
            add_inline_markdown(paragraph, token[1:-1], base_bold=base_bold,
                                base_italic=True, base_color=base_color)
        elif token.startswith("`"):
            _add_inline_code_run(paragraph, token[1:-1])
        else:  # link  [text](url)
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            link_text, url = lm.group(1), lm.group(2)
            _add_hyperlink(paragraph, link_text, url, bold=base_bold, italic=base_italic)
        pos = m.end()
    if pos < len(text):
        _plain_run(paragraph, text[pos:], base_bold, base_italic, base_color)


def _plain_run(paragraph, text, bold, italic, color):
    if text == "":
        return
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def _add_hyperlink(paragraph, text, url, bold=False, italic=False):
    """Render a link. A real http(s) URL becomes a clickable hyperlink (no noisy "(url)" tail).
    Sibling .md files and #anchors don't resolve in a standalone document, so they render as
    plain descriptive text — so "[Developer Guide](DEVELOPER_GUIDE.md)" reads as just
    "Developer Guide", not "Developer Guide (DEVELOPER_GUIDE.md)"."""
    if not (url.startswith("http://") or url.startswith("https://")):
        _plain_run(paragraph, text, bold, italic, None)
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    rfont = OxmlElement("w:rFonts")
    rfont.set(qn("w:ascii"), BODY_FONT)
    rfont.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(rfont)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --------------------------------------------------------------------------- #
# Block builders
# --------------------------------------------------------------------------- #
def add_heading(doc, level, text, first_content_heading):
    p = doc.add_paragraph()
    if level == 1 and not first_content_heading:
        # page break before every H1 except the first content heading
        _page_break_before(p)
    add_inline_markdown_heading(p, text, level)
    _keep_with_next(p)
    return p


def _page_break_before(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    pbb = OxmlElement("w:pageBreakBefore")
    ppr.append(pbb)


def add_inline_markdown_heading(paragraph, text, level):
    sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    colors = {1: DARK_BLUE, 2: DARK_BLUE, 3: LIGHT_BLUE, 4: DARK_BLUE}
    bold = {1: True, 2: True, 3: True, 4: False}
    italic = {1: False, 2: False, 3: False, 4: True}

    pf = paragraph.paragraph_format
    pf.space_before = Pt(16 if level <= 2 else 12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

    # headings may still contain inline code / bold; render token-aware, but
    # force the heading color/size/weight on plain segments.
    _render_heading_tokens(paragraph, text, sizes[level], colors[level], bold[level], italic[level])


def _render_heading_tokens(paragraph, text, size_pt, color, bold, italic):
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _heading_run(paragraph, text[pos:m.start()], size_pt, color, bold, italic, code=False)
        token = m.group(0)
        if token.startswith("**"):
            _heading_run(paragraph, token[2:-2], size_pt, color, True, italic, code=False)
        elif token.startswith("*"):
            _heading_run(paragraph, token[1:-1], size_pt, color, bold, True, code=False)
        elif token.startswith("`"):
            _heading_run(paragraph, token[1:-1], size_pt, color, bold, italic, code=True)
        else:
            lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            _heading_run(paragraph, lm.group(1), size_pt, color, bold, italic, code=False)
        pos = m.end()
    if pos < len(text):
        _heading_run(paragraph, text[pos:], size_pt, color, bold, italic, code=False)


def _heading_run(paragraph, text, size_pt, color, bold, italic, code):
    if text == "":
        return
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT if code else BODY_FONT
    run.font.size = Pt(size_pt - 1 if code else size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15
    add_inline_markdown(p, text)
    return p


def add_code_block(doc, lines):
    """Shaded, bordered code block. One paragraph per line keeps the box tight."""
    n = len(lines)
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(2) if i else Pt(6)
        pf.space_after = Pt(2) if i < n - 1 else Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.left_indent = Twips(120)
        pf.right_indent = Twips(120)
        _para_shading(p, CODE_FILL)
        # border only on outer edges so the block reads as one box
        sides = []
        sides.append("top") if i == 0 else None
        sides.append("bottom") if i == n - 1 else None
        sides += ["left", "right"]
        sides = [s for s in sides if s]
        _para_borders(p, CODE_BORDER, sz=4, space=4, sides=tuple(sides))
        run = p.add_run(line if line != "" else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
        # preserve leading whitespace
        run._r.get_or_add_rPr()
        t = run._r.find(qn("w:t"))
        if t is None:
            t = OxmlElement("w:t")
            run._r.append(t)
        t.set(qn("xml:space"), "preserve")


def add_callout(doc, lines):
    """Blockquote -> shaded bordered callout box (1x1 table for a true box)."""
    joined = " ".join(l.strip() for l in lines if l.strip())
    warn = bool(re.match(r"^\**\s*(warning|caution|important)", joined, re.IGNORECASE))
    fill = CALLOUT_RED if warn else CALLOUT_BLUE
    border = CALLOUT_BORDER_RED if warn else CALLOUT_BORDER_BLUE

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, _content_width_twips(doc))
    cell = table.cell(0, 0)
    _cell_shading(cell, fill)
    _cell_margins(cell, top=100, bottom=100, left=160, right=160)
    _set_table_borders(table, border)

    # clear default empty paragraph
    cell.paragraphs[0].text = ""
    first = True
    for raw in lines:
        stripped = raw.strip()
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.12
        # nested bullet inside a callout (e.g. "- Yes ...")
        m_bul = re.match(r"^[-*]\s+(.*)$", stripped)
        if m_bul:
            pf.left_indent = Inches(0.25)
            run = p.add_run("•  ")
            run.font.name = BODY_FONT
            add_inline_markdown(p, m_bul.group(1), base_italic=False)
        else:
            add_inline_markdown(p, stripped, base_italic=True)
    return table


def add_hr(doc):
    """Horizontal rule -> a thin spacer paragraph with a bottom border."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    _para_borders(p, TABLE_GRID, sz=6, space=1, sides=("bottom",))
    run = p.add_run("")
    run.font.size = Pt(2)


def add_list_item(doc, text, ordered, level, number=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.12
    pf.left_indent = Inches(0.30 + 0.30 * level)
    pf.first_line_indent = Inches(-0.22)
    if ordered:
        marker = f"{number}. "
    else:
        marker = "•  " if level == 0 else "◦  "
    mrun = p.add_run(marker)
    mrun.font.name = BODY_FONT
    if ordered:
        mrun.bold = True
    add_inline_markdown(p, text)
    return p


def add_table(doc, header_cells, body_rows):
    ncols = len(header_cells)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = _content_width_twips(doc)
    _set_table_width(table, total)
    _set_table_borders(table, TABLE_GRID)

    col_w = int(total / ncols)
    # header
    hdr = table.rows[0].cells
    for i, ctext in enumerate(header_cells):
        _cell_shading(hdr[i], "2E75B6")
        _cell_margins(hdr[i])
        _set_cell_width(hdr[i], col_w)
        cp = hdr[i].paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        _render_cell_content(cp, ctext, header=True)
    # body
    for r, row in enumerate(body_rows):
        cells = table.add_row().cells
        for i in range(ncols):
            val = row[i] if i < len(row) else ""
            _cell_shading(cells[i], "FFFFFF")
            _cell_margins(cells[i])
            _set_cell_width(cells[i], col_w)
            cp = cells[i].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.line_spacing = 1.05
            _render_cell_content(cp, val, header=False)
    return table


def _render_cell_content(paragraph, text, header):
    if header:
        # force white bold; still honor inline code font
        pos = 0
        for m in _INLINE_RE.finditer(text):
            if m.start() > pos:
                r = paragraph.add_run(text[pos:m.start()])
                _style_header_run(r)
            token = m.group(0)
            if token.startswith("`"):
                r = paragraph.add_run(token[1:-1])
                _style_header_run(r, code=True)
            elif token.startswith("**"):
                r = paragraph.add_run(token[2:-2])
                _style_header_run(r)
            elif token.startswith("*"):
                r = paragraph.add_run(token[1:-1])
                _style_header_run(r)
                r.italic = True
            else:
                lm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
                r = paragraph.add_run(lm.group(1))
                _style_header_run(r)
            pos = m.end()
        if pos < len(text):
            r = paragraph.add_run(text[pos:])
            _style_header_run(r)
    else:
        add_inline_markdown(paragraph, text)
        for r in paragraph.runs:
            if r.font.size is None:
                r.font.size = Pt(10)


def _style_header_run(run, code=False):
    run.font.name = CODE_FONT if code else BODY_FONT
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)


# --------------------------------------------------------------------------- #
# Width helpers
# --------------------------------------------------------------------------- #
def _content_width_twips(doc):
    """Usable content width, in EMU (page width minus left+right margins)."""
    sec = doc.sections[0]
    return int(sec.page_width - sec.left_margin - sec.right_margin)


def _set_table_width(table, width_twips):
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(int(width_twips / 635)))  # EMU->twips
    tblw.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_twips):
    cell.width = width_twips


# --------------------------------------------------------------------------- #
# Cover page
# --------------------------------------------------------------------------- #
def add_cover(doc, title, subtitle):
    # vertical spacing
    for _ in range(6):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.font.name = BODY_FONT
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    t.paragraph_format.space_after = Pt(6)

    # accent rule under the title
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_borders(rule, "2E75B6", sz=12, space=1, sides=("bottom",))
    rr = rule.add_run("")
    rr.font.size = Pt(2)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.name = BODY_FONT
    sr.font.size = Pt(18)
    sr.font.color.rgb = GRAY_SUB
    s.paragraph_format.space_before = Pt(10)

    for _ in range(2):
        doc.add_paragraph()

    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run("Azure DevOps Forager")
    mr.font.name = BODY_FONT
    mr.font.size = Pt(12)
    mr.bold = True
    mr.font.color.rgb = DARK_BLUE

    m2 = doc.add_paragraph()
    m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m2r = m2.add_run(f"Last updated: {LAST_UPDATED}")
    m2r.font.name = BODY_FONT
    m2r.font.size = Pt(11)
    m2r.font.color.rgb = GRAY_SUB

    # page break after the cover
    pb = doc.add_paragraph()
    pb.add_run().add_break()
    _page_break_before_hard(pb)


def _page_break_before_hard(paragraph):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# --------------------------------------------------------------------------- #
# Document base styling
# --------------------------------------------------------------------------- #
def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    # East-Asian font mapping so name sticks
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)

    # US Letter, 1" margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    return doc


# --------------------------------------------------------------------------- #
# Markdown parser / driver
# --------------------------------------------------------------------------- #
def parse_table_block(lines, start):
    """Return (header_cells, body_rows, next_index) for a pipe table at `start`."""
    def split_row(row):
        row = row.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [c.strip() for c in row.split("|")]

    header = split_row(lines[start])
    # lines[start+1] is the |---| separator
    body = []
    i = start + 2
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        body.append(split_row(lines[i]))
        i += 1
    return header, body, i


def is_table_start(lines, i):
    if "|" not in lines[i]:
        return False
    if i + 1 >= len(lines):
        return False
    sep = lines[i + 1].strip()
    return bool(re.match(r"^\|?[\s:|-]+\|[\s:|-]*$", sep)) and "-" in sep


def convert(md_path, out_path, title, subtitle):
    with open(md_path, "r", encoding="utf-8") as f:
        raw = f.read()
    # normalize line endings, drop the leading top-level '# Title' (goes on cover)
    lines = raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    doc = new_document()
    add_cover(doc, title, subtitle)

    i = 0
    first_content_heading = True
    skipped_title = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, block)
            continue

        # blockquote / callout (consume consecutive > lines)
        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                q = lines[i].strip()[1:]
                if q.startswith(" "):
                    q = q[1:]
                block.append(q)
                i += 1
            # drop trailing empty callout lines
            while block and block[-1].strip() == "":
                block.pop()
            add_callout(doc, block)
            continue

        # pipe table
        if is_table_start(lines, i):
            header, body, ni = parse_table_block(lines, i)
            add_table(doc, header, body)
            i = ni
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            htext = m.group(2).strip()
            if level == 1 and not skipped_title:
                # the document's top-level title -> already on cover, skip
                skipped_title = True
                i += 1
                continue
            add_heading(doc, level, htext, first_content_heading and level == 1)
            if level == 1:
                first_content_heading = False
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            add_hr(doc)
            i += 1
            continue

        # blank line
        if stripped == "":
            i += 1
            continue

        # ordered list item (supports simple nesting by leading spaces)
        m_ol = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ol:
            indent = len(m_ol.group(1))
            level = 1 if indent >= 3 else 0
            add_list_item(doc, m_ol.group(3).strip(), ordered=True,
                          level=level, number=m_ol.group(2))
            i += 1
            continue

        # unordered list item
        m_ul = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m_ul:
            indent = len(m_ul.group(1))
            level = 1 if indent >= 2 else 0
            add_list_item(doc, m_ul.group(2).strip(), ordered=False, level=level)
            i += 1
            continue

        # plain paragraph
        add_paragraph(doc, stripped)
        i += 1

    doc.save(out_path)
    return out_path


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main():
    jobs = [
        (
            os.path.join(DOCS_DIR, "USER_GUIDE.md"),
            os.path.join(DOCS_DIR, "AzureDevOpsForager_User_Guide.docx"),
            "Azure DevOps Forager",
            "User Guide",
        ),
        (
            os.path.join(DOCS_DIR, "DEVELOPER_GUIDE.md"),
            os.path.join(DOCS_DIR, "AzureDevOpsForager_Developer_Guide.docx"),
            "Azure DevOps Forager",
            "Developer Guide",
        ),
    ]
    for md, out, title, subtitle in jobs:
        if not os.path.exists(md):
            raise SystemExit(f"Source not found: {md}")
        convert(md, out, title, subtitle)
        size = os.path.getsize(out)
        print(f"[OK] {out}  ({size:,} bytes)")


if __name__ == "__main__":
    main()
